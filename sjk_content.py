# -*- encoding: utf-8 -*-
'''
@File    :   sjk_content.py
@Time    :   2022/04/21 11:16:09
@Author  :   Coder-Sakura
@Version :   1.0
@Desc    :   None
'''

# here put the import lib
import os
import re
import time
import random
from docx import Document
# 居中
from docx.enum.text import WD_ALIGN_PARAGRAPH
# (正文)修改字体
from docx.oxml.ns import qn
from docx.shared import Pt, Cm, RGBColor


from sjk_tool import logger, network_connect


class SJK_CONTENT:
	def remove_label(self, text_content):
		"""
		过滤文本
		:params text_content:待替换文本列表
		:return :list
		"""
		# 取出图片链接
		img_regular = re.compile(r""".*?src=['"](.*?)['"].*?""")
		# 去除转义字符,\n \a \t
		escape_regular = re.compile(r"[\a\b\f\n\t\r\v\\]")
		# 去除html标签
		html_regular = re.compile(r"<[^>]+>", re.S)
		# 去除utf编码保存文件导致出现的\ufeff
		bom_regular = re.compile(u"[\ufeff]+", re.S)

		res = []
		for text in text_content:
			# 匹配src链接
			if "src" in text and "img" in text:
				img = re.findall(img_regular, text.replace(" ", ""))
				for i in img:
					res.append(i)
			# 去除转义字符
			text = re.sub(escape_regular, "", text)
			# 去除html标签
			text = re.sub(html_regular, "", text)
			# 去除utf编码保存文件导致出现的\ufeff
			text = re.sub(bom_regular, "", text)
			if text != "":
				res.append(text)

		return res

	def find_size(self, index):
		size = []
		size_regular = re.compile(r""".*?width=['"](.*?)['"].*?height=['"](.*?)['"].*?""")
		try:
			size = re.findall(size_regular, self.text_content[index])
		except:
			return []
		else:
			if not size:
				return []
		# width height [('128', '36')] or []
		return size

	def get_img(self, url):
		img_path = os.path.join(
			os.path.dirname(self.docx_path),
			f"{str(int(time.time()))}-{str(random.choice(range(100)))}.jpg"
		)
		resp = network_connect(url)
		with open(img_path, "wb") as f:
			f.write(resp.content)

		return img_path

	def write_docx(self, section_name, docx_path, new_text_content):
		"""
		章节文档写入docx
		:params docx_path: docx存储路径
		:params new_text_content: 过滤过的内容
		:return : None
		"""
		document = Document()
		# 添加标题
		head = document.add_heading(level=1)
		head.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

		run  = head.add_run(section_name)
		run.font.name = u"微软雅黑"
		run.font.color.rgb = RGBColor(0,0,0)
		run._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')

		# 正文样式
		document.styles['Normal'].font.name = u'微软雅黑'
		document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')

		for k, _ in enumerate(new_text_content):
			# 插入图片
			if (_.startswith("http://") or _.startswith("https://")) and \
				any([".jpg" in _, ".png" in _, ".jpeg" in _]):
				try:
					img_path = self.get_img(_)
					shape = document.add_picture(img_path)
				except Exception as e:
					logger.warning(f"获取图片img链接 或 向docx内添加图片出错... - {_} - {e}")
					paragraph = document.add_paragraph(_, 'Normal')
					paragraph.paragraph_format.space_before = Pt(20)
				else:
					size = self.find_size(k)
					if size:
						pass
						# shape.width, shape.height = int(size[0][0]), int(size[0][1])
					else:
						raw_height, raw_width = shape.height, shape.width
						shape.height = Cm(3)
						shape.width = int(raw_width * shape.height / raw_height)
					# test
					os.remove(img_path)
			else:
				paragraph = document.add_paragraph(_, 'Normal')
				paragraph.paragraph_format.space_before = Pt(20)


		document.save(docx_path)
		logger.success(f"文档下载成功: {docx_path}")
		time.sleep(1)

	def main(self, section_name, docx_path, text_content):
		self.docx_path = docx_path
		self.text_content = text_content

		# 先正则过滤,再写入
		# 不存在则进行SC流程
		if os.path.exists(docx_path):
			logger.success(f"已存在: <docx_path> {docx_path}")
			return

		new_text_content = self.remove_label(text_content)
		self.write_docx(section_name, docx_path, new_text_content)